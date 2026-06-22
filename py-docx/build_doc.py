"""
Генерация DOCX из cleaned.json.

Архитектура (4 слоя):
  1. Чтение JSONL
  2. Нормализация текста (абзацы, переносы)
  3. Выбор пресета (preflight-оценка + postflight-валидация)
  4. Генерация DOCX (отдельная секция на каждое сообщение)

Корневые причины предыдущих дефектов:
  - Оценка по len(text) без учёта абзацев и длин строк
  - Одна секция на весь документ, меняющаяся на лету
  - Весь текст в одном абзаце — терялась структура
  - Justify на коротких строках создавал растяжку
"""

import argparse
import json
import math
import os
import sys

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Константы ────────────────────────────────────────────────

PT_TO_MM = 0.3528

PAGE_SIZES = {
    "A4": (210, 297),
    "A5": (148, 210),
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Пресеты от читаемого к компактному
COMPRESS_PRESETS = [
    {"font_size": 12, "margins_mm": 25, "line_spacing": 1.15, "before": 0, "after": 6},
    {"font_size": 11, "margins_mm": 22, "line_spacing": 1.15, "before": 0, "after": 4},
    {"font_size": 10, "margins_mm": 20, "line_spacing": 1.15, "before": 0, "after": 3},
    {"font_size": 10, "margins_mm": 17, "line_spacing": 1.1,  "before": 0, "after": 2},
    {"font_size": 9,  "margins_mm": 15, "line_spacing": 1.0,  "before": 0, "after": 2},
    {"font_size": 8,  "margins_mm": 12, "line_spacing": 1.0,  "before": 0, "after": 1},
]

# ── Слой 1: Чтение JSONL ────────────────────────────────────

def load_config(config_path):
    """Читает docx.json."""
    if not os.path.exists(config_path):
        return {}
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Генерация DOCX из cleaned.json",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры:
  python build_doc.py --input ../output --output ../output/book.docx
  python build_doc.py --input ../output --output ../output/book.docx --no-compress
        """,
    )
    parser.add_argument("--input", required=True, help="Папка с cleaned.json")
    parser.add_argument("--output", required=True, help="Путь к выходному .docx")
    parser.add_argument("--format", choices=["A4", "A5"], default=None, help="Формат страницы")
    parser.add_argument("--align", choices=["left", "center", "right", "justify"], default=None,
                        help="Выравнивание")
    parser.add_argument("--config", default=None, help="Путь к docx.json")
    parser.add_argument("--no-compress", action="store_true",
                        help="Фиксированный формат без сжатия")
    return parser.parse_args()


def resolve_settings(args):
    cfg = {}
    if args.config:
        cfg = load_config(args.config)
    return {
        "page_format": args.format or cfg.get("page_format", "A4"),
        "alignment": args.align or cfg.get("alignment", "left"),
        "font_name": cfg.get("font_name", "Times New Roman"),
        "font_size": cfg.get("font_size", 12),
        "no_compress": args.no_compress or cfg.get("no_compress", False),
    }


def read_messages(json_path):
    """Читает JSONL, возвращает список текстов."""
    if not os.path.exists(json_path):
        print(f"ОШИБКА: Файл не найден: {json_path}")
        sys.exit(1)

    messages = []
    with open(json_path, "r", encoding="utf-8") as f:
        for line_num, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                text = obj.get("message", "")
                if text:
                    messages.append(text)
            except json.JSONDecodeError as e:
                print(f"Предупреждение: строка {line_num} — невалидный JSON: {e}")
    return messages

# ── Слой 2: Нормализация текста ──────────────────────────────

def normalize_text(text):
    """
    Нормализует текст:
    - Сохраняет структуру абзацев (пустые строки = разделитель)
    - Убирает тройные+ пустые строки, оставляет двойные
    - Возвращает список абзацев (каждый абзац — строка)
    """
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    paragraphs = []
    current = []

    for line in lines:
        stripped = line.rstrip()
        if stripped == "":
            # Пустая строка — конец текущего абзаца
            if current:
                paragraphs.append(" ".join(current))
                current = []
            # Добавляем разделитель абзацев (не более одного)
            if paragraphs and paragraphs[-1] != "":
                paragraphs.append("")
        else:
            current.append(stripped)

    if current:
        paragraphs.append(" ".join(current))

    # Убираем ведущие/замыкающие пустые абзацы
    while paragraphs and paragraphs[0] == "":
        paragraphs.pop(0)
    while paragraphs and paragraphs[-1] == "":
        paragraphs.pop()

    return paragraphs if paragraphs else [""]

# ── Слой 3: Оценка и выбор пресета ──────────────────────────

def usable_width_mm(page_width_mm, margins_mm):
    """Ширина текстовой области в мм."""
    return page_width_mm - 2 * margins_mm


def usable_height_mm(page_height_mm, margins_mm):
    """Высота текстовой области в мм."""
    return page_height_mm - 2 * margins_mm


def line_height_mm(font_size_pt, line_spacing):
    """Высота одной строки в мм."""
    return font_size_pt * PT_TO_MM * line_spacing


def max_chars_per_line(font_size_pt, page_width_mm, margins_mm):
    """
    Максимальное количество символов в строке.
    Использует консервативную оценку: средняя ширина символа = 0.6 * размер
    (для кириллицы и пропорциональных шрифтов это безопасная граница).
    """
    uw = usable_width_mm(page_width_mm, margins_mm)
    avg_char_w = font_size_pt * 0.6 * PT_TO_MM
    if avg_char_w <= 0:
        avg_char_w = 1.0
    return max(1, int(uw / avg_char_w))


def max_lines_on_page(font_size_pt, line_spacing, page_height_mm, margins_mm):
    """Сколько строк помещается по высоте."""
    uh = usable_height_mm(page_height_mm, margins_mm)
    lh = line_height_mm(font_size_pt, line_spacing)
    return max(1, int(uh / lh))


def preflight_estimate(paragraphs, font_size_pt, page_w, page_h, margins_mm,
                       line_spacing, space_before, space_after, alignment):
    """
    Preflight-оценка: считает общее количество строк с учётом:
    - абзацев и отступов между ними
    - длины каждой строки
    - межстрочного интервала
    - вертикальных отступов абзацев
    """
    max_chars = max_chars_per_line(font_size_pt, page_w, margins_mm)
    max_lines = max_lines_on_page(font_size_pt, line_spacing, page_h, margins_mm)
    lh = line_height_mm(font_size_pt, line_spacing)

    total_lines = 0

    for idx, para_text in enumerate(paragraphs):
        if para_text == "":
            # Пустой абзац-разделитель: занимает ~0.5 строки
            total_lines += 0.5
            continue

        # Количество строк в абзаце
        para_len = len(para_text)
        if para_len == 0:
            continue

        lines_in_para = math.ceil(para_len / max_chars)
        total_lines += lines_in_para

        # Вертикальные отступы абзаца (space_before + space_after)
        para_height = space_before + space_after
        total_lines += para_height / lh if lh > 0 else 0

    return total_lines, max_lines


def postflight_check(paragraphs, font_size_pt, page_w, page_h, margins_mm,
                     line_spacing, space_before, space_after):
    """
    Postflight-проверка: более строгая оценка с учётом:
    - самой длинной строки в тексте (не средняя)
    - реальных переносов слов
    - запаса 5% на непредвиденное
    """
    max_chars = max_chars_per_line(font_size_pt, page_w, margins_mm)
    max_lines = max_lines_on_page(font_size_pt, line_spacing, page_h, margins_mm)
    lh = line_height_mm(font_size_pt, line_spacing)

    total_height_mm = 0

    for idx, para_text in enumerate(paragraphs):
        if para_text == "":
            total_height_mm += lh * 0.5
            continue

        if not para_text:
            continue

        # Считаем строки, разбивая по пробелам для учёта переносов слов
        words = para_text.split()
        if not words:
            continue

        current_line_len = 0
        lines = 1
        for word in words:
            word_len = len(word) + (1 if current_line_len > 0 else 0)
            if current_line_len + word_len > max_chars and current_line_len > 0:
                lines += 1
                current_line_len = len(word)
            else:
                current_line_len += word_len

        para_height = lines * lh + space_before + space_after
        total_height_mm += para_height

    usable_h = usable_height_mm(page_h, margins_mm)

    # Запас 5%
    return total_height_mm <= usable_h * 0.95


def find_best_preset(paragraphs, page_w, page_h, alignment):
    """
    Последовательно перебирает пресеты.
    Для каждого: preflight → postflight.
    Возвращает (preset, overflow).
    """
    for preset in COMPRESS_PRESETS:
        total_lines, max_lines = preflight_estimate(
            paragraphs, preset["font_size"], page_w, page_h,
            preset["margins_mm"], preset["line_spacing"],
            preset["before"], preset["after"], alignment
        )

        fits_preflight = total_lines <= max_lines

        if fits_preflight:
            # Postflight-проверка
            fits_postflight = postflight_check(
                paragraphs, preset["font_size"], page_w, page_h,
                preset["margins_mm"], preset["line_spacing"],
                preset["before"], preset["after"]
            )
            if fits_postflight:
                return preset, False

    # Ни один пресет не прошёл проверку
    return COMPRESS_PRESETS[-1], True

# ── Слой 4: Генерация DOCX ──────────────────────────────────

def new_section(doc, page_w_mm, page_h_mm, margins_mm):
    """
    Добавляет новую секцию с заданными параметрами страницы.
    python-docx: add_section() принимает start_type.
    """
    new_sec = doc.add_section()
    new_sec.page_width = Mm(page_w_mm)
    new_sec.page_height = Mm(page_h_mm)
    new_sec.top_margin = Mm(margins_mm)
    new_sec.bottom_margin = Mm(margins_mm)
    new_sec.left_margin = Mm(margins_mm)
    new_sec.right_margin = Mm(margins_mm)
    return new_sec


def apply_preset_to_paragraph(paragraph, preset, font_name, alignment):
    """Применяет пресет к абзацу."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(preset["before"])
    pf.space_after = Pt(preset["after"])
    pf.line_spacing = preset["line_spacing"]

    if alignment in ALIGNMENTS:
        paragraph.alignment = ALIGNMENTS[alignment]

    for run in paragraph.runs:
        run.font.size = Pt(preset["font_size"])
        run.font.name = font_name


def add_page_break(doc):
    """Добавляет разрыв страницы через XML."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)
    # Убираем отступы у параграфа с break'ом
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0


def create_document(messages, settings):
    """
    Генерирует DOCX: каждое сообщение — отдельная секция (страница).
    Возвращает (документ, stats).
    """
    doc = Document()
    page_w, page_h = PAGE_SIZES[settings["page_format"]]
    alignment = settings["alignment"]
    font_name = settings["font_name"]

    # Первая секция (дефолтная) — удаляем лишний параграф
    first_section = doc.sections[0]
    first_section.page_width = Mm(page_w)
    first_section.page_height = Mm(page_h)

    stats = {"total": len(messages), "fitted": 0, "overflow": 0, "warnings": []}

    for i, text in enumerate(messages):
        # Нормализуем текст в абзацы
        paragraphs = normalize_text(text)

        # Пропускаем пустые сообщения
        if all(p == "" for p in paragraphs):
            print(f"  Сообщение {i+1}: пустое, пропущено")
            continue

        # Выбираем пресет
        if settings["no_compress"]:
            preset = {
                "font_size": settings["font_size"],
                "margins_mm": 25,
                "line_spacing": 1.15,
                "before": 0,
                "after": 6,
            }
            overflow = False
        else:
            preset, overflow = find_best_preset(
                paragraphs, page_w, page_h, alignment
            )

        # Добавляем страницу: break перед сообщением (кроме первого)
        if i > 0:
            add_page_break(doc)

        # Новая секция для каждого сообщения
        section = new_section(doc, page_w, page_h, preset["margins_mm"])

        # Добавляем абзацы
        for para_text in paragraphs:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(para_text)
            run.font.name = font_name
            run.font.size = Pt(preset["font_size"])
            apply_preset_to_paragraph(paragraph, preset, font_name, alignment)

        # Лог
        total_lines, max_lines = preflight_estimate(
            paragraphs, preset["font_size"], page_w, page_h,
            preset["margins_mm"], preset["line_spacing"],
            preset["before"], preset["after"], alignment
        )

        if overflow:
            stats["overflow"] += 1
            stats["warnings"].append(i + 1)
            print(f"  [{i+1}] пресет: {preset['font_size']}pt/{preset['margins_mm']}mm "
                  f"| строк: ~{total_lines:.0f}/{max_lines} "
                  f"| WARNING: не поместилось, минимальный формат")
        else:
            stats["fitted"] += 1
            print(f"  [{i+1}] пресет: {preset['font_size']}pt/{preset['margins_mm']}mm "
                  f"| строк: ~{total_lines:.0f}/{max_lines} OK")

    return doc, stats


def main():
    args = parse_args()
    settings = resolve_settings(args)

    json_path = os.path.join(args.input, "cleaned.json")
    print(f"Чтение: {json_path}")
    messages = read_messages(json_path)
    print(f"Найдено сообщений: {len(messages)}")

    if not messages:
        print("Нет сообщений для обработки")
        sys.exit(0)

    print(f"Формат: {settings['page_format']}, выравнивание: {settings['alignment']}")
    print(f"Шрифт: {settings['font_name']}, размер: {settings['font_size']}pt")
    if settings["no_compress"]:
        print("Режим: фиксированный (без сжатия)")
    print()

    doc, stats = create_document(messages, settings)

    output_dir = os.path.dirname(args.output)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc.save(args.output)

    print()
    print(f"Обработано: {stats['total']} сообщений")
    print(f"  Уместилось: {stats['fitted']}")
    if stats["overflow"] > 0:
        print(f"  Не поместилось: {stats['overflow']} (минимальный формат)")
        print(f"  Предупреждения для сообщений: {stats['warnings']}")
    print(f"Результат: {args.output}")


if __name__ == "__main__":
    main()
